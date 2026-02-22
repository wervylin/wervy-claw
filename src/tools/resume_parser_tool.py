"""
简历解析工具 - 用于解析Word/PDF格式的简历文件
"""
import os
from typing import Optional
from pypdf import PdfReader
from docx import Document
from langchain.tools import tool


@tool
def parse_resume_file(file_path: str) -> str:
    """
    解析简历文件（支持Word和PDF格式），提取文本内容

    Args:
        file_path: 简历文件的绝对路径，例如 "/tmp/resume.docx" 或 "/tmp/resume.pdf"

    Returns:
        str: 提取的简历文本内容

    注意：
        - 文件路径必须是绝对路径
        - 支持的文件格式：.docx, .pdf
        - 如果文件不存在或格式不支持，将返回错误信息
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return f"错误：文件 '{file_path}' 不存在，请检查文件路径是否正确。"

        # 获取文件扩展名
        file_ext = os.path.splitext(file_path)[1].lower()

        # 根据文件扩展名选择解析方式
        if file_ext == '.pdf':
            return _parse_pdf(file_path)
        elif file_ext == '.docx':
            return _parse_docx(file_path)
        else:
            return f"错误：不支持的文件格式 '{file_ext}'，仅支持 .docx 和 .pdf 格式。"

    except Exception as e:
        return f"简历文件解析失败: {str(e)}"


def _parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    try:
        reader = PdfReader(file_path)
        text_content = []
        
        # 提取每一页的文本
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        
        # 合并所有页面文本
        full_text = "\n\n".join(text_content)
        
        if not full_text.strip():
            return "警告：PDF文件中未提取到文本内容，可能是扫描版PDF，请尝试手动输入或提供可编辑的文件。"
        
        return f"简历内容（PDF）：\n\n{full_text}"
        
    except Exception as e:
        raise Exception(f"PDF解析失败: {str(e)}")


def _parse_docx(file_path: str) -> str:
    """解析Word (.docx) 文件"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # 合并所有文本
        full_text = "\n".join(text_content)
        
        if not full_text.strip():
            return "警告：Word文件中未提取到文本内容，请检查文件是否为空。"
        
        return f"简历内容（Word）：\n\n{full_text}"
        
    except Exception as e:
        raise Exception(f"Word文件解析失败: {str(e)}")


@tool
def validate_resume_text(resume_text: str) -> str:
    """
    验证简历文本是否符合基本要求

    Args:
        resume_text: 简历文本内容

    Returns:
        str: 验证结果，包含文本长度信息和完整性评估
    """
    try:
        # 去除空白字符后计算长度
        clean_text = resume_text.strip()
        text_length = len(clean_text)
        
        # 评估文本长度
        if text_length < 100:
            return f"⚠️ 简历内容过短（仅{text_length}字），建议补充更多详细信息，包括：教育背景、工作经历、技能、项目经验等。"
        elif text_length > 10000:
            return f"⚠️ 简历内容过长（{text_length}字），建议精简内容，突出重点，控制在1-2页（约1000-3000字）之间。"
        else:
            return f"✅ 简历文本长度适中（{text_length}字），可以进行分析。"
            
    except Exception as e:
        return f"简历文本验证失败: {str(e)}"
